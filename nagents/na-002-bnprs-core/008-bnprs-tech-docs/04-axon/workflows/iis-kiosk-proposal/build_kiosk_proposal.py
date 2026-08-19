#!/usr/bin/env python3
"""Next-Gen Self-Service Instant Card Issuance Kiosk — Full Proposal (generic / white-label)."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

# BNPRS brand theme: dark green (#2D4A3E) primary + gold (#D4952B) accent.
NAVY = RGBColor(0x2D, 0x4A, 0x3E); NAVY_HEX = "2D4A3E"   # primary (brand green)
GOLD = RGBColor(0xB0, 0x7A, 0x1E); GOLD_HEX = "D4952B"   # accent gold (text uses a darker shade for contrast)
RED = RGBColor(0xC0, 0x00, 0x00); GREY = RGBColor(0x59, 0x59, 0x59)
GREEN = RGBColor(0x2D, 0x4A, 0x3E); ZEBRA = "EFF4F0"     # light green-tint zebra
MEDGREEN = RGBColor(0x3E, 0x7D, 0x5A)                     # medium green (K3 uses 2E75B6 for H3/accent)
CREAM = "F3EAD7"                                          # light gold/cream fill for metadata keys
HSIZE = {1: 18, 2: 15, 3: 13}                            # K3 exact heading sizes (pt)
HCOLOR = {1: NAVY, 2: NAVY, 3: MEDGREEN}                 # K3 H1/H2 navy(green), H3 blue(med-green)
DOC_ID = "PROP-ICISKIOSK-2026-001"
IMG = "/Users/bnprs/BPR/GitRepos1/aim.pat/nagents/na-002-bnprs-core/008-bnprs-tech-docs/04-axon/workflows/iis-kiosk-proposal/kiosk_imgs"
IMG_M = "/Users/bnprs/BPR/GitRepos1/aim.pat/nagents/na-002-bnprs-core/008-bnprs-tech-docs/04-axon/workflows/iis-kiosk-proposal/mutek"

doc = Document()
st = doc.styles["Normal"]; st.font.name = "Arial"; st.font.size = Pt(11)   # K3: Arial 11pt
sec = doc.sections[0]
from docx.shared import Inches as _In
sec.page_width = _In(8.5); sec.page_height = _In(11)                        # K3: US Letter
sec.top_margin = Cm(2.54); sec.bottom_margin = Cm(2.54)                     # K3: 1 inch all sides
sec.left_margin = Cm(2.54); sec.right_margin = Cm(2.54)

def shade(cell, hexcolor):
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(sh)
def cell_text(cell, text, bold=False, color=None, size=10):
    cell.text = ""; p = cell.paragraphs[0]
    for part in re.split(r"(MENTA)", str(text)):   # client name always bold
        if not part: continue
        r = p.add_run(part); r.font.size = Pt(size); r.bold = bold or part == "MENTA"
        if part == "MENTA": r.font.color.rgb = NAVY
        elif color is not None: r.font.color.rgb = color
def _hrule(p):  # thin bottom border under a heading (K3 H1 pBdr)
    pPr = p._p.get_or_add_pPr(); b = OxmlElement("w:pBdr"); bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6"); bot.set(qn("w:space"), "4"); bot.set(qn("w:color"), GOLD_HEX)
    b.append(bot); pPr.append(b)
def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = "Arial"; r.font.size = Pt(HSIZE.get(level, 13)); r.bold = True
        r.font.color.rgb = HCOLOR.get(level, NAVY)
    if level == 1: _hrule(h)
    return h
def para(text, italic=False, color=None, size=11, bold=False):
    p = doc.add_paragraph()
    for part in re.split(r"(MENTA)", text):        # client name always bold
        if not part: continue
        r = p.add_run(part); r.italic = italic; r.font.size = Pt(size); r.bold = bold or part == "MENTA"
        if part == "MENTA": r.font.color.rgb = NAVY
        elif color is not None: r.font.color.rgb = color
    return p
def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True; r.font.color.rgb = NAVY; p.add_run(text)
    else: p.add_run(text)
def note(text):
    p = doc.add_paragraph(); r = p.add_run(text); r.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = GREY
def table(headers, rows, widths=None, zebra=True, hsize=10):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; shade(c, NAVY_HEX); cell_text(c, h, bold=True, color=RGBColor(0xFF,0xFF,0xFF), size=hsize)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cell_text(cells[i], val)
            if zebra and ri % 2 == 0: shade(cells[i], ZEBRA)
    if widths:
        for i, w in enumerate(widths):
            for r in t.rows: r.cells[i].width = w
    return t
def figure(path, caption, width=Cm(16.0)):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=width)
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run(caption); r.italic = True; r.font.size = Pt(9); r.font.color.rgb = GREY
def pb(): doc.add_page_break()
def toc():
    p = doc.add_paragraph(); run = p.add_run()
    fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"), "begin")
    ins = OxmlElement("w:instrText"); ins.set(qn("xml:space"), "preserve"); ins.text = 'TOC \\o "1-2" \\h \\z \\u'
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    tt = OxmlElement("w:t"); tt.text = "Right-click → Update Field to build the Table of Contents."
    en = OxmlElement("w:fldChar"); en.set(qn("w:fldCharType"), "end")
    for e in (fb, ins, sep, tt, en): run._r.append(e)

# header / footer
hp = sec.header.paragraphs[0]; hp.text = ""
r = hp.add_run(f"{DOC_ID}  |  "); r.font.size = Pt(8.5); r.font.color.rgb = GREY
r2 = hp.add_run("CONFIDENTIAL"); r2.font.size = Pt(8.5); r2.bold = True; r2.font.color.rgb = RED
r3 = hp.add_run("\t\tInstant Card Issuance Solution (Kiosk) — Proposal"); r3.font.size = Pt(8.5); r3.font.color.rgb = GREY
fp = sec.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT; fp.text = ""
rr = fp.add_run("Page "); rr.font.size = Pt(8.5); rr.font.color.rgb = GREY
fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE"); fp._p.append(fld)

# ── COVER ── (vertically centred on page 1 by the PDF pipeline)
for txt, sz, col, bold in [("SOLUTION DESIGN & COMMERCIAL PROPOSAL", 22, NAVY, True),
                           ("Next-Generation Self-Service", 15, NAVY, True),
                           ("Instant Card Issuance Solution (Kiosk)", 15, NAVY, True)]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt); r.bold = bold; r.font.size = Pt(sz); r.font.color.rgb = col
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("A modular, secure, production-ready self-service platform for instant, on-demand personalisation and dispensing of payment and ID cards.")
r.font.size = Pt(11); r.font.color.rgb = GREY; r.italic = True
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("CONFIDENTIAL"); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = RED
doc.add_paragraph()
meta = [("Document ID", DOC_ID), ("Version", "1.0 (High-Level Proposal)"), ("Classification", "CONFIDENTIAL"),
        ("Date", "August 2026"), ("Prepared By", "BNPRS"),
        ("Manufacturing Partner", "Alpha91 KP Solutions (enclosure fabrication)"),
        ("Prepared For", "MENTA — Managing Director review"), ("Status", "For approval — multiple design options")]
mt = doc.add_table(rows=0, cols=2); mt.style = "Table Grid"; mt.alignment = WD_TABLE_ALIGNMENT.CENTER
for k, v in meta:
    cells = mt.add_row().cells
    shade(cells[0], CREAM); cell_text(cells[0], k, bold=True, color=NAVY, size=10)
    cell_text(cells[1], v, color=(RED if k == "Classification" else None), bold=(k == "Classification"), size=10)
    cells[0].width = Cm(4.8); cells[1].width = Cm(11)
pb()

# ── REVISION HISTORY ── (K3 house style: sits after cover, before the TOC)
heading("Revision History", 1)
table(["Version", "Date", "Author", "Description"],
      [["1.0", "18-Aug-2026", "BNPRS", "Initial high-level proposal — design concept, Basic / Pro / Max design options, "
        "per-model general-arrangement drawings, hardware/electrical/control design and indicative commercials."]],
      widths=[Cm(2.2), Cm(2.6), Cm(3.4), Cm(8.0)])
pb()

# ── EXECUTIVE SUMMARY ──
heading("Executive Summary", 1)
para("This proposal presents the design and commercial basis for a next-generation self-service kiosk that "
     "personalises and dispenses payment and ID cards to a customer instantly, on demand, with no staff "
     "intervention. The objective is not to build another kiosk, but to establish a complete, scalable, "
     "production-ready platform that sets a new benchmark for self-service instant card issuance — and can "
     "grow into a flagship product line for the region and international markets.")
para("The solution is engineered in-house: our team leads the full electronic system design, software, and "
     "integration. Our mechanical partner is responsible solely "
     "for fabricating the enclosure to our engineering specifications. The industrial design, electrical "
     "architecture, and control system have already been developed to a working 3D model and schematic level, "
     "de-risking the path to prototype.", )
para("Key differentiators:", bold=True, color=NAVY)
for lead, txt in [("Tri-modal biometrics — ", "facial, iris and fingerprint capture for strong identity assurance."),
                  ("Super-User mode — ", "a privileged operating mode for authorised staff to issue cards on behalf of customers, enabling centralised preparation and secure distribution."),
                  ("In-machine card embossing — ", "beyond flat personalisation, supporting premium and legacy card formats (optional module)."),
                  ("Modular vendor panel — ", "a standardised, quick-release work-surface panel lets any component vendor be swapped in minutes without modifying the kiosk body."),
                  ("Vendor-agnostic — ", "the design is not tied to any single supplier or brand: every subsystem (biometrics, MRZ, card engine/embosser, printer, controller) is selected against a standard interface, so components can be sourced from competing vendors on price, availability or certification — avoiding vendor lock-in. The MUTEK engine in §6.1, for example, is one option, not a fixed dependency."),
                  ("Defence-grade data handling — ", "multi-layer encryption (mTLS end-to-end + HSM at rest), logical and physical security compliance, and cryptographic obfuscation of rejected/defective cards."),
                  ("Resilient dual-controller architecture — ", "the industrial PC runs the UI and all peripherals while a dedicated Raspberry Pi controller handles chip encoding exclusively, isolating the secure encoding path and adding an independent monitoring/watchdog layer.")]:
    bullet(txt, bold_lead=lead)
para("Three design options (Basic / Pro / Max) are offered so the capability and investment level can be "
     "matched to the deployment strategy.")
pb()

# ── TOC ── (a real Table of Contents is generated for the PDF via pandoc --toc)

# ── 1. INTRODUCTION ──
heading("1. Introduction", 1)
heading("1.1 Purpose", 2)
para("This document defines the solution concept, design options, technical architecture, industrial design, "
     "security model, operational workflows, delivery approach and indicative commercials for a self-service "
     "instant card issuance kiosk. It is intended as a high-level proposal for management approval, offering "
     "multiple design options for selection.")
heading("1.2 Scope", 2)
for b in ["The end-to-end self-service card-issuance journey and a privileged Super-User issuance mode",
          "Kiosk hardware (personalisation/printing/embossing, dispensing, biometrics, customer interface, security module, controllers, power)",
          "Software (kiosk application, personalisation engine, Super-User service, host/CMS integration, fleet management)",
          "Industrial design and ergonomics of the enclosure",
          "Security, key management and compliance",
          "Indicative commercials with optional-feature cost impact (embossing and deposit modules priced both with and without)"]:
    bullet(b)
heading("1.3 Definitions and Abbreviations", 2)
table(["Term", "Definition"],
      [["CMS", "Card Management System — the host that authorises and drives personalisation"],
       ["EMV", "Chip-card standard for payment cards (contact and contactless)"],
       ["HSM", "Hardware Security Module — external key-protection service (Cloud HSM or on-premise server HSM); the kiosk connects to it and does not contain one"],
       ["KYC", "Know Your Customer — identity verification"],
       ["mTLS", "Mutual TLS — two-way authenticated, encrypted transport"],
       ["MRZ", "Machine-Readable Zone — data band on passports/ID documents"],
       ["EPP", "Encrypting PIN Pad"],
       ["BOM / CAPEX / OPEX", "Bill of Materials / capital expenditure / operating expenditure"]],
      widths=[Cm(3.2), Cm(12.8)])
pb()

# ── 2. UNDERSTANDING OF REQUIREMENTS ──
heading("2. Understanding of Requirements", 1)
para("The following requirements have been captured and are addressed in this proposal.")
table(["#", "Requirement", "Addressed in"],
      [["R-01", "Self-service instant card issuance and dispensing", "§3, §5, §6"],
       ["R-02", "Multi-biometric authentication (facial, iris, fingerprint)", "§5.2"],
       ["R-03", "Super-User mode — issue cards on behalf of customers; centralised preparation & secure distribution", "§5.3, §10.2"],
       ["R-04", "Card embossing capability (optional module)", "§5.4, §6, §12"],
       ["R-05", "Multi-layer data encryption (mTLS + HSM)", "§9.1"],
       ["R-06", "Logical and physical security compliance", "§9"],
       ["R-07", "Data obfuscation for rejected/defective cards", "§9.3"],
       ["R-08", "Deposit / acceptance module (optional)", "§5.5, §6, §12"],
       ["R-09", "Support for all card types & schemes (contact, contactless, magstripe, emboss)", "§5.4"],
       ["R-10", "Host/CMS integration and identity verification via partner application", "§7.3"],
       ["R-11", "Multiple design options for selection", "§4"]],
      widths=[Cm(1.3), Cm(11.2), Cm(3.5)])
pb()

# ── 3. SOLUTION CONCEPT ──
heading("3. Solution Concept & Vision", 1)
heading("3.1 Concept", 2)
para("The kiosk is an ATM-style, free-standing self-service terminal for indoor deployment in bank branches. "
     "A customer authenticates, confirms their details, and requests a card. The kiosk authorises the request "
     "with the bank host, personalises a blank card inside the machine (chip encoding, magstripe, print and — "
     "optionally — embossing), verifies it, and dispenses the finished card within minutes. The same platform "
     "additionally supports a privileged Super-User mode for staff-assisted and centralised issuance.")
heading("3.2 Journeys", 2)
para("Self-service journey:", bold=True, color=NAVY)
for i, s in enumerate(["Customer selects 'Issue Card' on the touchscreen",
                       "Identity verification — document (MRZ), biometrics (face/iris/fingerprint), and the partner identity app",
                       "Details confirmed; request authorised with the host/CMS",
                       "Card picked from a hopper and personalised (chip + magstripe + print (+ emboss))",
                       "Card verified in-machine; rejected cards are data-obfuscated and binned",
                       "Card dispensed; receipt issued; audit record written"], 1):
    bullet(f"Step {i}: {s}")
para("Super-User journey:", bold=True, color=NAVY)
for b in ["Authorised operator authenticates into the privileged mode",
          "Issues one or a batch of cards on behalf of customers (centralised preparation)",
          "Cards are securely held/tracked for later distribution to cardholders",
          "All privileged actions are strongly authenticated and fully audited"]:
    bullet(b)
pb()

# ── 4. DESIGN OPTIONS ──
heading("4. Design Options", 1)
para("Three configuration options — Basic, Pro and Max — are offered on a single common kiosk body, work "
     "surface and software platform. They are not three separate products but three capability tiers of one "
     "design: the enclosure, the modular vendor panel, the dual-controller architecture and the card-production "
     "pipeline are shared across all three, and a model is defined simply by which subsystems and software "
     "modules are populated on that common base.")
para("This common-platform approach gives the buyer direct commercial and operational advantages:", bold=True, color=NAVY)
for lead, txt in [("Field-upgradable — ", "a site can start at Basic and move up to Pro or Max later by adding modules (iris scanner, embosser, deposit unit) and enabling the corresponding software — without replacing the kiosk body or re-fabricating the enclosure."),
                  ("Common spares & servicing — ", "one body, one vendor-panel interface and one software build serve every tier, minimising spare-part inventory, technician training and integration effort across a mixed fleet."),
                  ("Vendor-agnostic sourcing — ", "because every subsystem mounts to a standard interface, the components that distinguish the tiers can be sourced from competing vendors on price, lead-time or certification (see §8.1)."),
                  ("Predictable cost curve — ", "capability is added incrementally, so per-unit investment tracks the value each tier unlocks — indicative pricing per option is given in §12.")]:
    bullet(txt, bold_lead=lead)
para("The three tiers map directly to deployment intent. Basic suits high-volume, cost-sensitive rollouts where "
     "unattended instant issuance with dual-modal biometrics (face + fingerprint) and baseline security is "
     "sufficient. Pro — the recommended option — adds tri-modal identity assurance (face + iris + fingerprint), "
     "privileged Super-User issuance for staff-assisted and centralised distribution, hardware-backed key "
     "protection (HSM), reject-card data obfuscation and standard card embossing; it delivers every "
     "capability identified as essential for a bank-branch deployment. Max is the flagship tier: on top of Pro it "
     "adds the cash / document deposit module, scheme-level certification and advanced fleet management and "
     "analytics, for the most demanding, high-throughput or regulated sites. The comparison below summarises how "
     "capability builds up across the three tiers.")
table(["Capability", "Basic", "Pro (Recommended)", "Max"],
      [["Self-service instant issuance", "Yes", "Yes", "Yes"],
       ["Biometrics", "Face + Fingerprint", "Face + Iris + Fingerprint", "Face + Iris + Fingerprint"],
       ["MRZ / passport + barcode capture", "Yes", "Yes", "Yes"],
       ["Super-User mode", "—", "Yes", "Yes"],
       ["Card embossing", "—", "Standard", "Standard"],
       ["Card hoppers (multiple card types)", "2", "up to 6", "up to 6 + replenish"],
       ["Security (mTLS + HSM, logical+physical)", "Baseline", "Full", "Full + scheme certified"],
       ["Reject-card data obfuscation", "—", "Yes", "Yes"],
       ["Deposit / acceptance module", "—", "—", "Standard"],
       ["Remote fleet management & analytics", "Basic", "Standard", "Advanced"],
       ["Dual-controller resilience", "Yes", "Yes", "Yes + redundancy"]],
      widths=[Cm(5.3), Cm(3.0), Cm(4.4), Cm(3.3)], hsize=9)
note("Card embossing is standard on Pro and Max; the cash/document deposit module is standard on Max. "
     "Indicative per-unit pricing for each option is given in Section 12 (Commercials).")

heading("4.1 Model Drawings & Dimensions", 2)
para("The three models share a common platform, work surface and modular vendor panel; they differ in "
     "capability set and, for Max, in enclosure size. External dimensions and standard additions are "
     "summarised below, with the certified engineering drawing of each model on the following pages.")
table(["Model", "W × D × H (mm)", "Standard additions vs Basic", "Drawing No."],
      [["Basic", "600 × 700 × 1600", "—", "KIOSK-V2-600-002"],
       ["Pro", "600 × 700 × 1600", "Iris, Super-User, HSM, card embossing", "KIOSK-V2-600-002"],
       ["Max", "650 × 700 × 1700", "Pro + deposit module, scheme-certified, fleet mgmt", "KIOSK-V3-650-003"]],
      widths=[Cm(2.2), Cm(4.0), Cm(6.9), Cm(3.5)], hsize=9)
pb()

heading("Basic model", 3)
para("The entry configuration: self-service instant issuance with dual-modal biometrics (facial + "
     "fingerprint), MRZ/passport and barcode capture, and baseline security — in a compact "
     "600 × 700 × 1600 mm enclosure on a common modular vendor panel.")
figure(f"{IMG}/variant_basic.png", "Figure 4.1 — Basic model — general arrangement (side, front, top views).", width=Cm(16.5))
pb()

heading("Pro model (Recommended)", 3)
para("The recommended configuration: adds iris biometrics (tri-modal identity assurance), Super-User mode, "
     "HSM-backed security, reject-card data obfuscation and standard card embossing — in the same "
     "600 × 700 × 1600 mm footprint as Basic.")
figure(f"{IMG}/variant_pro.png", "Figure 4.2 — Pro model — general arrangement (side, front, top views).", width=Cm(16.5))
pb()

heading("Max model", 3)
para("The flagship configuration: everything in Pro plus the cash/document deposit module, scheme "
     "certification and advanced fleet management — in a larger 650 × 700 × 1700 mm enclosure sized to "
     "house the deposit module with service clearance.")
figure(f"{IMG}/variant_max.png", "Figure 4.3 — Max model — general arrangement (side, front, top views).", width=Cm(16.5))
pb()

# ── 5. FUNCTIONAL CAPABILITIES ──
heading("5. Functional Capabilities", 1)
heading("5.1 Instant Card Personalisation & Dispensing", 2)
para("Blank cards are held in card hoppers (multiple hoppers support multiple card products). The integrated "
     "card personalisation module encodes the chip, writes the magstripe, prints the card face and — where the "
     "embossing module is fitted — embosses the card, before verification and dispensing. All card types are "
     "supported: contact and contactless (dual-interface) EMV, magstripe, and embossed cards.")
heading("5.2 Multi-Biometric Authentication", 2)
para("The kiosk performs tri-modal biometric capture for strong identity assurance: facial image (camera), "
     "iris (iris scanner) and fingerprint (fingerprint scanner), complemented by document capture via the "
     "MRZ/passport reader and barcode/QR reader.")
heading("5.3 Super-User Mode", 2)
para("A software-controlled privileged mode enables authorised personnel to issue cards on behalf of "
     "customers. This extends the platform beyond self-service to operational scenarios such as centralised "
     "card preparation and secure distribution to cardholders. Access is strongly authenticated and every "
     "action is audited.")
heading("5.4 Card Embossing", 2)
para("Card embossing adds raised-character personalisation for premium and legacy card formats. It integrates "
     "into the same card-production pipeline and power architecture, and is standard on the Pro and Max models.")
heading("5.5 Cash / Document Deposit Module", 2)
para("The deposit/acceptance module extends the kiosk to accept cash, documents or media as part of onboarding "
     "workflows, and is standard on the Max model.")
pb()

# ── 6. HARDWARE, ELECTRICAL & CONTROL DESIGN ──
heading("6. Hardware, Electrical & Control Design", 1)

heading("6.1 Card Personalisation & Embossing Engine", 2)
para("At the heart of the kiosk is the MUTEK MTK-F8-1400-DYC18 Turntable Card Personalising Dispenser — an "
     "integrated module that stores, personalises, embosses and dispenses cards in a single card path. Card "
     "embossing is built into this engine, so no separate embosser is required: the module performs chip "
     "read/write, magstripe encoding, embossing/indenting and tipping, and returns any defective card to a "
     "dedicated recycle bin. The kiosk uses pre-printed card stock and delivers instant electrical "
     "personalisation and raised-character embossing. This integrated engine is one option for providing the "
     "card embossing that is standard on the Pro and Max models (see §4 and §5.4).")
figure(f"{IMG_M}/mutek-views.png", "Figure 6.1 — MUTEK MTK-F8-1400-DYC18 personalisation & embossing engine — six orthographic views. Visible: card library turntable, six card boxes, transport rail and the personalisation / embossing head.", width=Cm(16.5))
table(["Parameter", "Value"],
      [["Personalisation", "Chip R/W, magstripe encoding, Embossing/Indenting (51 digits), Tipping"],
       ["Card library (turntable)", "Dispenses any card in the library; 100% accuracy"],
       ["Card boxes (hoppers)", "6 boxes × 100 cards = 600 blanks"],
       ["Recycle / reject bin", "2 boxes × 120 cards — defective cards stored, re-issued or data-obfuscated"],
       ["Card issuing speed", "≤ 60 s (issue + emboss/indent + tipping)"],
       ["Card size", "85.5 × 54 mm, 0.76–1.1 mm thick"],
       ["Interface / Power", "RS-232 · DC 24 V ±5%, 5 A (idle 700 mA)"],
       ["Environment", "0–50 °C, 0–90% RH (non-condensing)"]],
      widths=[Cm(5.0), Cm(11.0)])

heading("6.2 Subsystem Overview", 2)
table(["Subsystem", "Function"],
      [["MUTEK card personalisation / embossing engine", "Chip write, magstripe, print, emboss/indent, tipping; turntable dispense; reject bin"],
       ["Industrial PC (main controller)", "Kiosk application, UI and personalisation engine"],
       ["Raspberry Pi 5 controller", "Handles chip encoding exclusively (dedicated secure encoding controller); independent monitoring & watchdog"],
       ["Customer interface", "Touch display, MRZ/passport reader, barcode reader, camera, iris scanner, fingerprint scanner, encrypting PIN pad"],
       ["Secure key access", "Encrypted channel to an external HSM (Cloud or on-premise server) for key storage & cryptographic operations — no HSM inside the kiosk"],
       ["Power", "AC → 1 kVA UPS → dual-output SMPS → 24 V (card engine/printer), 12 V (PC/aux), 5 V (RPi)"],
       ["Receipt printer, status LED, fans", "Receipt issuance, status indication, thermal management"]],
      widths=[Cm(5.5), Cm(10.5)])

heading("6.3 Electrical & Control Architecture", 2)
para("The electrical and control system has been engineered to full schematic level across four sheets — AC "
     "power distribution, DC power distribution, control & data interface, and protective earth — confirming "
     "feasibility ahead of prototyping. Each sheet is shown below with an explanation of how that part of the "
     "circuit works.")

heading("6.3.1 AC Power Distribution (KSK-001)", 3)
figure(f"{IMG}/ksk001_ac.png", "Figure 6.2.1 — AC power distribution.", width=Cm(13.5))
para("AC mains (110/230 VAC, 50/60 Hz) enters at terminal block X1 as Live, Neutral and Protective Earth. A "
     "two-pole main power switch (QS1) isolates the whole kiosk, and a 10 A MCB/fuse (F1) protects the live line "
     "only. Power then passes through a 1 kVA online UPS (UPS1) — so the kiosk keeps running and can complete a "
     "card already in production through a mains dip — and on to the dual-output SMPS (PS1). Protective earth is "
     "carried straight through to the earth bus and is deliberately not bonded to neutral.")

heading("6.3.2 DC Power Distribution (KSK-002)", 3)
figure(f"{IMG}/ksk002_dc.png", "Figure 6.2.2 — DC power distribution.", width=Cm(13.5))
para("The dual-output SMPS provides two DC rails. The 24 V rail (fused distribution TB1) powers the two "
     "highest-load subsystems — the MUTEK card personalising/embossing engine and the thermal printer. The 12 V "
     "rail (TB2) powers the industrial PC and the auxiliary I/O board, and feeds a 12→5 V DC/DC converter that "
     "powers the Raspberry Pi controller. Every output is individually fused (F2–F8, ratings set from measured "
     "load) so a fault in one subsystem cannot bring down the others.")

heading("6.3.3 Control & Data Interface (KSK-003)", 3)
figure(f"{IMG}/ksk003_ctrl.png", "Figure 6.2.3 — Control & data interface.", width=Cm(13.5))
para("The industrial PC (PC1) runs the kiosk application and drives the UI and all peripherals — the touch "
     "monitor (HDMI), MRZ/passport scanner, fingerprint scanner, barcode reader, cameras, PCI encrypting PIN "
     "pad, thermal printer, the card print/emboss & dispensing engine, and the deposit module. A dedicated "
     "Raspberry Pi (RPI1) is responsible for chip encoding exclusively — isolating the secure "
     "chip-personalisation path — and provides an independent monitoring and watchdog layer separate from the "
     "main PC.")

heading("6.3.4 Protective Earth (KSK-004)", 3)
figure(f"{IMG}/ksk004_earth.png", "Figure 6.2.4 — Protective earth.", width=Cm(13.5))
para("A single protective-earth bus bar bonds every chassis — UPS, SMPS, PC, printer, auxiliary I/O board, "
     "Raspberry Pi, the card engine, the deposit module and the kiosk frame — to a common low-impedance earth. "
     "Earth is used for safety only and is never used as a DC return.")
note("The main controller may be provisioned as a Windows industrial PC or an Android industrial panel PC per "
     "final component selection; the schematic shows the Android panel-PC arrangement and the architecture "
     "supports either.")
pb()

# ── 7. SOFTWARE ARCHITECTURE ──
heading("7. Software Architecture", 1)
table(["Component", "Role"],
      [["Kiosk Application", "Customer UI, workflow orchestration, localisation, accessibility"],
       ["Personalisation Engine", "Drives the card module — magstripe encoding, print, emboss, verification (chip encoding handled by the dedicated RPi controller)"],
       ["Super-User Service", "Privileged issuance workflows, strong authentication, audit"],
       ["Device Manager", "Abstracts and drives peripherals (via the modular vendor panel)"],
       ["Host / CMS Connector", "Authorises requests and exchanges card data with the bank's CMS"],
       ["Identity Integration", "Customer identity verification via the partner identity application"],
       ["Security Services", "mTLS, key management (HSM), reject-card obfuscation"],
       ["Monitoring & Fleet Management", "Health, consumables, remote management and analytics"]],
      widths=[Cm(5.3), Cm(10.7)])
heading("7.3 Host / CMS Integration", 2)
para("The Card Management System is provided by MENTA; the kiosk integrates to it through a secure "
     "connector. Customer identity verification is performed through the partner identity application, and all "
     "traffic is protected with mutual TLS.")
pb()

# ── 8. INDUSTRIAL DESIGN & UX ──
heading("8. Industrial Design & UX", 1)
para("The industrial design has been developed to a finished 3D model with defined ergonomics. The free-standing "
     "enclosure presents a clean, secure face with a top biometric module (camera + iris), a tilted touch "
     "display, an angled interaction panel carrying the readers and PIN pad, a receipt/card outlet, and a "
     "lockable equipment cabinet on lockable casters with levelling feet.")
figure(f"{IMG}/img2.png", "Figure 8.1 — Optimised industrial design: 3D views and the modular, interchangeable vendor-panel concept.", width=Cm(13.5))
heading("8.1 Modular Vendor Panel — a platform differentiator", 2)
para("The interaction readers (MRZ/OCR, fingerprint, barcode, PIN pad) are mounted on a standardised, "
     "quick-release vendor panel. A common mounting interface (standard size, hole pattern, connectors and "
     "cable routing) lets an entire vendor combination be swapped in minutes without modifying the kiosk body. "
     "This makes the platform genuinely vendor-agnostic: no subsystem is locked to a single supplier, and "
     "components can be sourced from competing vendors on price, availability or certification.")
for lead, txt in [("Vendor flexibility — ", "support different component vendors with a panel change."),
                  ("Quick replacement — ", "swap or upgrade in minutes; the kiosk body is unchanged."),
                  ("Standardised platform — ", "one body serves all configurations and options."),
                  ("Future-ready — ", "adopt new technologies via new panels.")]:
    bullet(txt, bold_lead=lead)
heading("8.2 Ergonomics & Dimensions", 2)
table(["Feature", "Height from floor (mm)"],
      [["Cabinet / wheel datum", "80"], ["Receipt + card outlet (centreline)", "820"],
       ["Interaction panel (front / centre / rear)", "926 / 970 / 1014"],
       ["Display (bottom / centre / top)", "1080 / 1227.5 / 1375"],
       ["Iris centre", "1490"], ["Camera centre", "1570"], ["Overall top", "1600"]],
      widths=[Cm(8.5), Cm(7.5)])
table(["Parameter", "Value"],
      [["Interaction panel angle", "25° from horizontal"], ["Display angle", "10° from vertical"],
       ["Overall dimensions (W × D × H)", "600 × 700 × 1600 mm"]],
      widths=[Cm(8.5), Cm(7.5)])
figure(f"{IMG}/kiosk-final.png", "Figure 8.2 — Finished 3D model of the kiosk (mechanical partner, Alpha91 KP Solutions) — synchronised with the industrial-design study above.", width=Cm(6.0))
heading("8.3 Design Coordination — Enclosure and Internal Modules", 2)
para("The BNPRS industrial-design study (Figure 8.1) and the mechanical partner's finished 3D model "
     "(Figure 8.2) represent the same, synchronised design — overall 600 × 700 × 1600 mm with the ergonomics of "
     "§8.2. During detailed engineering the internal module layout is verified against this common model. The "
     "principal coordination item is the card engine: the MUTEK module footprint (≈ 665 × 536 mm) fits within "
     "the 600 × 700 mm cabinet when oriented with its 665 mm dimension running front-to-back (within the 700 mm "
     "depth) and its 536 mm dimension across the width (within the 600 mm width), leaving service clearance. The "
     "following are confirmed against the model before design freeze:")
for b in ["Overall external dimensions and internal cabinet volume (in agreement across both figures)",
          "Card-engine orientation, mounting and service access (door swing, card-box loading)",
          "Card and receipt outlet positions (outlet centreline at 820 mm) aligned to the engine's dispense path",
          "Interaction-panel cut-outs matching the modular vendor-panel interface",
          "Cable routing, earthing points and ventilation/thermal paths"]:
    bullet(b)
note("Enclosure fabrication is performed by the mechanical partner (Alpha91 KP Solutions) to BNPRS engineering "
     "specifications; all electronics, integration and software are delivered by BNPRS.")
pb()

# ── 9. SECURITY & COMPLIANCE ──
heading("9. Security & Compliance", 1)
heading("9.1 Multi-Layer Encryption", 2)
para("All communication with the host/CMS and back-end services uses mutual TLS (mTLS) for two-way "
     "authenticated, encrypted transport. Cryptographic keys are held and used in an external Hardware Security "
     "Module — a Cloud HSM or an on-premise server HSM (the kiosk itself does not contain an HSM); sensitive "
     "data at rest is encrypted. Keys are referenced by Key Version and Key Check Value (KCV) — key values "
     "are never stored in documentation or configuration.")
heading("9.2 Logical & Physical Security", 2)
for b in ["Strong authentication and role-based access (customer, Super-User, technician, administrator)",
          "Lockable cabinet, tamper indication, and secured casters",
          "Full audit logging of issuance and privileged actions",
          "Alignment with PCI, EMV and card-scheme issuance requirements, and applicable data-protection law"]:
    bullet(b)
heading("9.3 Rejected / Defective Card Data Obfuscation", 2)
para("Any card that fails in-machine verification, or is otherwise rejected/defective, has its written data "
     "cryptographically obfuscated (overwritten/zeroised) before it is moved to the reject bin — ensuring no "
     "readable cardholder or key data can leave the machine on a discarded card.")

heading("9.4 Component-wise Certifications for PCI Compliance", 2)
para("PCI compliance is achieved at system level (a PCI DSS / PIN-Security programme), but each payment-path "
     "component must itself carry the appropriate certification. The table maps every relevant component to the "
     "certification it should hold. Selecting pre-certified components is the fastest route to an overall "
     "compliant machine.")
table(["Component", "Required / Recommended Certification", "Purpose"],
      [["Encrypting PIN Pad (EPP)", "PCI PTS POI (current version)", "Secure PIN entry & tamper protection — mandatory"],
       ["Chip card reader (in card engine)", "EMV Level 1 & Level 2 (EMVCo)", "Contact chip acceptance/personalisation"],
       ["Contactless reader (if fitted)", "EMV Contactless L1 + scheme (Visa/Mastercard)", "Contactless card acceptance"],
       ["External HSM (Cloud / on-prem server)", "FIPS 140-2/3 Level 3 + PCI HSM", "Cryptographic key generation, storage & use — outside the kiosk"],
       ["Kiosk & payment software", "PCI DSS + PCI Secure Software (SSF) / P2PE", "Cardholder-data security in software"],
       ["Card personalisation / print", "Card-scheme personalisation approval (Visa/Mastercard/RuPay; Qi Card)", "Compliant card issuance & data prep"],
       ["End-to-end transport", "TLS 1.2+/mTLS; P2PE (if adopted)", "Encrypted, authenticated communications"],
       ["Whole system & operations", "PCI DSS, PCI PIN Security", "Overall compliance of the deployed machine"],
       ["Electrical / hardware", "CE, FCC, RoHS (+ regional safety/EMC)", "Product safety & EMC (prerequisite to sell)"],
       ["Biometric modules", "ISO/IEC 19794 (+ regional accreditation as applicable)", "Biometric data interoperability & quality"]],
      widths=[Cm(4.6), Cm(6.6), Cm(4.8)], hsize=9)
note("Certification is a cost and lead-time driver: pre-certified modules (especially the PCI-PTS PIN pad, the "
     "PCI-HSM/FIPS security module and EMVCo-approved reader) should be prioritised in procurement to minimise "
     "the compliance programme's scope and duration.")
pb()

# ── 10. OPERATIONAL WORKFLOWS ──
heading("10. Operational Workflows", 1)
heading("10.1 Self-Service Issuance", 2)
para("The standard unattended flow described in §3.2, from customer request to dispensed card, with automatic "
     "exception handling for jams, personalisation failures, empty hoppers and host timeouts.")
heading("10.2 Super-User / Centralised Issuance", 2)
para("Authorised staff issue single or batch cards on behalf of customers for centralised preparation and "
     "secure distribution — extending the platform to operational and branch-back-office scenarios while "
     "preserving strong authentication and full audit.")
heading("10.3 Maintenance", 2)
para("Consumable replenishment (blank cards, ribbon), preventive maintenance and remote monitoring; component "
     "service or upgrade is simplified by the modular vendor panel.")
pb()

# ── 11. DELIVERY ROADMAP ──
heading("11. Delivery Roadmap", 1)
table(["Phase", "Activity", "Outcome"],
      [["1. Approval", "Approve option & scope", "Go-ahead"],
       ["2. Detailed Engineering", "Finalise hardware/software design & BOM", "Design freeze"],
       ["3. Procurement", "Source components; engage enclosure partner", "Parts on hand"],
       ["4. Build & Integration", "Assemble & integrate first unit", "Prototype"],
       ["5. Validation", "Functional, security & issuance testing", "Validated unit"],
       ["6. Manufacturing Readiness", "Pilot batch, documentation", "Production-ready"]],
      widths=[Cm(4.5), Cm(6.5), Cm(5.0)])
para("Following approval and completion of component procurement, first-unit delivery is targeted within "
     "45 days.", italic=True)
pb()

# ── 12. COMMERCIALS ──
heading("12. Commercials", 1)
para("Pricing is in USD and preliminary, prepared for budgeting; most component prices are expected within "
     "±5% of the estimates once supplier quotations are finalised. Figures below are the current hardware "
     "Bill of Materials; engineering, software and margin are quoted separately. Optional modules (embossing "
     "and deposit) are shown as add-on deltas so the impact of each feature on cost is explicit.")
heading("12.1 Component-wise Hardware Cost (Base Configuration)", 2)
bom = [["1", "Fanless Box PC (12th Gen i7, 16 GB, Win 11)", "1", "705"],
       ["2", "Barcode / QR Code Scanner", "1", "73"],
       ["3", "Fingerprint Scanner", "1", "104"],
       ["4", "Passport / MRZ Scanner", "1", "835"],
       ["5", "Encrypting PIN Pad", "1", "251"],
       ["6", "Card Printer", "1", "991"],
       ["7", "Iris Scanner", "1", "57"],
       ["8", "Industrial SMPS", "1", "157"],
       ["9", "Smart UPS (1000 VA)", "1", "261"],
       ["10", "Receipt Printer", "1", "115"],
       ["11", "Web Camera (facial)", "1", "73"],
       ["12", "21.5\" Touch Display", "1", "261"],
       ["13", "Custom Kiosk Body (enclosure)", "1", "1305"],
       ["14", "Raspberry Pi 5 (8 GB) controller", "1", "94"],
       ["15", "Industrial 64 GB microSD", "1", "13"],
       ["16", "Raspberry Pi Power Adapter", "1", "13"],
       ["17", "Card Dispenser Hopper (500-card)", "6", "2664"],
       ["", "Base Hardware Subtotal (USD)", "", "7,972"]]
table(["#", "Component", "Qty", "Cost (USD)"], bom, widths=[Cm(1.2), Cm(9.3), Cm(1.5), Cm(4.0)])
note("Base configuration excludes the optional embossing and deposit modules (see below). Cryptographic key "
     "protection uses an external Cloud or on-premise server HSM, provided separately — it is not part of the "
     "kiosk supply and is not priced here.")

heading("12.2 Optional Modules — Cost Impact (With / Without)", 2)
table(["Optional Module", "Included in", "Incremental Cost (USD)", "Notes"],
      [["Card Embossing unit", "Pro / Max", "+ 3,500", "Raised-character personalisation (emboss/indent + tipping)"],
       ["Cash / document deposit interface", "Max", "+ 3,000 – 5,000", "Deposit/acceptance module (range by configuration)"],
       ["External HSM (Cloud / on-premise server)", "Pro / Max", "Not included", "Key protection provided outside the kiosk (customer-supplied Cloud or on-prem server HSM); see §9.1"],
       ["Iris upgrade (adds iris to Basic)", "included in Pro/Max", "+ « preliminary »", "Tri-modal biometrics"],
       ["Additional hoppers (beyond base)", "as required", "≈ 444 / hopper", "500-card capacity each"]],
      widths=[Cm(4.6), Cm(2.6), Cm(4.4), Cm(4.4)], hsize=9)

heading("12.3 Indicative Per-Unit Price by Option", 2)
table(["Option", "Configuration", "Indicative Price Band / unit (USD)"],
      [["Basic", "Base, dual biometric, baseline security", "« lower band »"],
       ["Pro (Recommended)", "Pro base + tri-modal + Super-User + card embossing (external HSM integration)", "« mid band »"],
       ["Max", "Pro + deposit + certified + fleet mgmt", "« upper band »"]],
      widths=[Cm(4.6), Cm(7.4), Cm(4.0)])
note("Target all-in per-unit range for the intended configuration: USD 8,000–16,000, subject to final "
     "component quotations, optional modules selected, engineering, software and margin (quoted separately). "
     "All prices exclusive of taxes and duties.")
pb()

# ── 13. TEAM & PARTNERSHIP ──
heading("13. Team & Partnership Model", 1)
table(["Party", "Responsibility"],
      [["MENTA", "Provides the Card Management System (CMS) and identity application (SuperQi); business requirements and approvals"],
       ["BNPRS", "Complete electronic system design, software, integration, testing and delivery"],
       ["Alpha91 KP Solutions", "Fabrication and assembly of the kiosk enclosure to BNPRS engineering specifications only"]],
      widths=[Cm(4.5), Cm(11.5)])
pb()

# ── 14. RISKS, ASSUMPTIONS & INDEPENDENCE ──
heading("14. Risks & Assumptions", 1)
heading("14.1 Assumptions", 2)
for b in ["Pricing is preliminary and for budgeting; most components within ±5%",
          "Card personalisation module selection to be confirmed (candidate identified)",
          "CMS and identity application are provided by MENTA",
          "Indoor deployment in bank branches; standard throughput per module capacity"]:
    bullet(b)
pb()

# ── 15. APPENDICES ──
heading("15. Appendices", 1)
heading("15.1 Appendix A — Full Bill of Materials", 2)
note("See §12.1. Grand total (base hardware): USD 7,972 (preliminary, ±5%).")
heading("15.2 Appendix B — Design Reference", 2)
note("Industrial-design renders (Figures 8.1–8.2), ergonomic dimensions (§8.2) and engineered electrical/"
     "control schematics (Figure 6.1) are available in full resolution.")
heading("15.3 Appendix C — Glossary & Vendor Shortlist", 2)
note("Extended glossary and vendor shortlist available on request (key values never embedded).")
e = doc.add_paragraph(); e.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = e.add_run("— End of Document —"); r.italic = True; r.font.color.rgb = GREY

OUT = "/Users/bnprs/BPR/GitRepos1/aim.pat/nagents/na-002-bnprs-core/008-bnprs-tech-docs/07-axon-terminals/deliverables/Instant-Card-Issuance-Kiosk-Proposal.docx"
doc.save(OUT); print("saved", OUT)
